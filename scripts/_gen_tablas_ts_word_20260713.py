"""
Tablas de Búsqueda Tabú (TS) en formato Word (.docx), para pegar/insertar
directamente en un documento.

Mismos datos que las tablas LaTeX de _gen_tablas_mh_20260713.py:
  1. Resultados por instancia con parámetros de la corrida ganadora
     (87 filas; costo comparable con ajuste delta CARPLIB en val; costo en
     negritas cuando iguala la BKS).
  2. Soluciones explícitas como arcos (u,v) — rutas separadas por '|',
     depósito implícito — en sección apaisada con márgenes de 1 cm.

Salida:
    resultados/ts_tablas_word_20260713.docx

Uso:
    python scripts/_gen_tablas_ts_word_20260713.py
"""
from __future__ import annotations

import statistics
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt

RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAIZ / "scripts"))
from run_val_egl_20260710 import INSTANCIAS_VAL_GDB, INSTANCIAS_EGL  # noqa: E402
from _gen_tabla_cuckoo_20260712 import (  # noqa: E402
    INSTANCIAS_SMALL, _deltas, _mapas_tr, _solucion_arcos,
)
from _gen_tablas_mh_20260713 import MHS, _mejores  # noqa: E402

DESTINO = RAIZ / "resultados" / "ts_tablas_word_20260713.docx"

CAPTION_RES = (
    "Table 1. Tabu Search: best result per instance and parameters of the "
    "winning run. Costs follow the CARPLIB convention (the constant delta "
    "adjustment is applied on the val family); costs matching the BKS are "
    "shown in bold. θ is the tabu tenure and B the number of random "
    "neighbors sampled and evaluated per iteration. cfg: P = solo_p_inter, "
    "B = binario_capacidad, R = pr_aislado (Path Relinking) on the "
    "23-instance calibration set; T = val/egl transfer campaign "
    "(class-tuned configuration, pr_aislado, budget of 10^6 evaluations / "
    "300 s per run).")

CAPTION_SOL = (
    "Table 2. Tabu Search: best solution found per instance (the same "
    "winning runs as Table 1). Costs are reported on the BKS reference "
    "scale; costs matching the BKS are shown in bold. Each solution is "
    "listed as an array of required edges (u,v): routes are separated by "
    "'|' and the depot is implicit at both ends of every route.")


def _celda(cel, texto: str, *, bold: bool = False, size: float = 8.0,
           center: bool = False) -> None:
    p = cel.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if center:
        p.alignment = WD_TABLE_ALIGNMENT.CENTER
    r = p.add_run(texto)
    r.bold = bold
    r.font.size = Pt(size)


def _caption(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(8)
    r.italic = True


def main() -> None:
    spec = MHS["ts"]
    deltas = _deltas()
    mapas = _mapas_tr()
    mejores = _mejores(spec)
    orden = INSTANCIAS_SMALL + INSTANCIAS_VAL_GDB + INSTANCIAS_EGL
    filas = [(i, mejores[i]) for i in orden if i in mejores]

    doc = Document()
    # Sección 1 (vertical): tabla de resultados.
    doc.add_heading("Tabu Search — results and solutions per instance", 1)
    _caption(doc, CAPTION_RES)

    enc = ["Instance", "BKS", "Cost", "Gap %", "t (s)", "θ", "B",
           "p_inter / cfg"]
    t1 = doc.add_table(rows=1, cols=len(enc))
    t1.style = "Table Grid"
    for j, txt in enumerate(enc):
        _celda(t1.rows[0].cells[j], txt, bold=True, center=True)

    gaps = []
    for inst, b in filas:
        costo = b["_costo"] + deltas.get(inst, 0.0)
        bks = float(b["bks_referencia"])
        gap = (costo - bks) / bks * 100.0
        gaps.append(gap)
        fila = t1.add_row().cells
        _celda(fila[0], inst)
        _celda(fila[1], f"{bks:.0f}")
        _celda(fila[2], f"{costo:.0f}", bold=costo <= bks)
        _celda(fila[3], f"{gap:.2f}")
        _celda(fila[4], f"{b['_tiempo']:.0f}")
        _celda(fila[5], f"{float(b['tabu_tenure']):.0f}")
        _celda(fila[6], f"{float(b['tam_vecindario']):.0f}")
        _celda(fila[7], f"{b.get('p_inter') or '--'} ({b['_origen']})")
    resumen = t1.add_row().cells
    n_bks = sum(1 for g in gaps if g <= 0.0001)
    _celda(resumen[0], f"Mean gap: {statistics.mean(gaps):.2f}%", bold=True)
    _celda(resumen[2], f"BKS reached: {n_bks}/{len(gaps)}", bold=True)

    # Sección 2 (apaisada, márgenes 1 cm): tabla de soluciones.
    sec = doc.add_section()
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1))

    _caption(doc, CAPTION_SOL)
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    for j, txt in enumerate(("Instance", "Cost", "Solution")):
        _celda(t2.rows[0].cells[j], txt, bold=True, center=True)
    t2.columns[0].width = Cm(2.6)
    t2.columns[1].width = Cm(1.9)
    t2.columns[2].width = Cm(23.2)

    for inst, b in filas:
        costo = b["_costo"] + deltas.get(inst, 0.0)
        bks = float(b["bks_referencia"])
        sol = _solucion_arcos(b.get("mejor_solucion_tr_legible", ""),
                              mapas.get(inst, {}))
        sol = sol.replace("),(", "), (")
        fila = t2.add_row().cells
        _celda(fila[0], inst, size=7)
        _celda(fila[1], f"{costo:.0f}", bold=costo <= bks, size=7)
        _celda(fila[2], sol, size=6)
        # Fija el ancho por celda (Word ignora columns.width en tablas
        # autoajustables si no se replica en cada fila).
        for cel, w in zip(fila, (2.6, 1.9, 23.2)):
            cel.width = Cm(w)

    doc.save(DESTINO)
    print(f"Filas: {len(filas)}  ->  {DESTINO}")


if __name__ == "__main__":
    main()
